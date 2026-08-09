from pathlib import Path
from docx import Document
from openpyxl import Workbook

BAHAN = Path("tests/bahan")

d = Document()
d.add_heading("Notulen Rapat Pleno Semester Ganjil", level=1)
d.add_paragraph("SDN Sukamaju, Kecamatan Sumedang Selatan")
d.add_paragraph(
    "Kepala sekolah menugaskan wakil kurikulum menyusun jadwal supervisi "
    "akademik untuk enam guru kelas pada minggu ketiga."
)
d.add_paragraph(
    "Bendahara melaporkan penyerapan dana BOS tahap pertama sebesar delapan "
    "puluh dua persen, dengan sisa dialihkan ke pemeliharaan ruang kelas."
)
d.add_paragraph(
    "Nomor kepegawaian narasumber tercatat sebagai NIP 199901019999019999 "
    "dan telepon panitia 081200000000."
)
d.save(BAHAN / "notulen.docx")

w = Workbook()
s = w.active
s.title = "Serapan"
s.append(["Uraian", "Pagu", "Realisasi", "Sisa"])
s.append(["Honor guru", 12000000, 9800000, "=B2-C2"])
s.append(["Pemeliharaan", 5000000, 3100000, "=B3-C3"])
w.save(BAHAN / "serapan.xlsx")
print("docx dan xlsx dibuat")
