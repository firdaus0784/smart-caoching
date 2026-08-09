"""DOCX berperubahan terlacak — bahan uji B-3.

Disusun dengan menyisipkan unsur w:ins dan w:del langsung ke XML, karena
python-docx tidak memiliki antarmuka untuk perubahan terlacak. Isinya dibuat
agar ketiga keadaan muncul pada satu paragraf: teks biasa, teks disisipkan,
dan teks dihapus.
"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

BAHAN = Path("tests/bahan")
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _run(teks: str, tag: str = "t") -> etree._Element:
    r = etree.SubElement(etree.Element(qn("w:wrap")), qn("w:r"))
    t = etree.SubElement(r, qn(f"w:{tag}"))
    t.text = teks
    t.set(qn("xml:space"), "preserve")
    return r


d = Document()
p = d.add_paragraph()
p._p.append(_run("Jadwal supervisi disusun untuk "))

ins = etree.SubElement(p._p, qn("w:ins"))
ins.set(qn("w:id"), "1")
ins.set(qn("w:author"), "Wakil Kurikulum")
ins.set(qn("w:date"), "2026-08-01T00:00:00Z")
ins.append(_run("enam"))

p._p.append(_run(" guru kelas, "))

dele = etree.SubElement(p._p, qn("w:del"))
dele.set(qn("w:id"), "2")
dele.set(qn("w:author"), "Wakil Kurikulum")
dele.set(qn("w:date"), "2026-08-01T00:00:00Z")
dele.append(_run("BATAL DIHAPUS", tag="delText"))

p._p.append(_run("bukan delapan."))

d.save(BAHAN / "notulen-terlacak.docx")
print("dibuat:", (BAHAN / "notulen-terlacak.docx").stat().st_size, "bita")

print(
    "paragraph.text bawaan python-docx ->",
    repr(Document(str(BAHAN / "notulen-terlacak.docx")).paragraphs[0].text),
)

# Paragraf kedua: penghapusan yang memakai w:t, bukan w:delText.
# Sebagian penghasil DOCX menulisnya begini, dan pengekstrak yang hanya
# melewati w:delText akan memasukkan teks yang sudah dihapus.
d2 = Document(str(BAHAN / "notulen-terlacak.docx"))
p2 = d2.add_paragraph()
p2._p.append(_run("Anggaran disetujui "))
del2 = etree.SubElement(p2._p, qn("w:del"))
del2.set(qn("w:id"), "3")
del2.set(qn("w:author"), "Bendahara")
del2.set(qn("w:date"), "2026-08-01T00:00:00Z")
del2.append(_run("SALAH TERHAPUS", tag="t"))
p2._p.append(_run("sepenuhnya."))
d2.save(BAHAN / "notulen-terlacak.docx")
print("paragraf kedua ditambahkan")
